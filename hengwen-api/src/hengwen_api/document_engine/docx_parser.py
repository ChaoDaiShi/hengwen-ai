import re
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as load_docx
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from hengwen_api.core.exceptions import AppError, ErrorCode
from hengwen_api.document_engine.models import (
    DocumentModel,
    FigureModel,
    HeadingModel,
    ParagraphModel,
    ReferenceModel,
    RunModel,
    SectionModel,
    TableModel,
)

REFERENCE_HEADING = re.compile(r"^参考文献\s*$")
REFERENCE_ENTRY = re.compile(r"^\s*(?:\[(\d+)\]|(\d+)[.、])\s*(.+)$")
CAPTION_PATTERN = re.compile(r"^\s*([图表])\s*([0-9]+(?:[-.][0-9]+)*)\s*(.*)$")


def _invalid_document(exc: Exception | None = None) -> AppError:
    return AppError(
        ErrorCode.INVALID_DOCUMENT,
        "无法解析该 DOCX 文档",
        status_code=422,
    )


def _run_font_name(run: Run) -> str | None:
    font = run.font
    if font.name:
        return str(font.name)
    properties = run._element.rPr
    if properties is None or properties.rFonts is None:
        return None
    east_asia = properties.rFonts.get(qn("w:eastAsia"))
    return str(east_asia) if east_asia else None


def _line_spacing(paragraph: Paragraph) -> float | None:
    spacing = paragraph.paragraph_format.line_spacing
    if spacing is None:
        return None
    if hasattr(spacing, "pt"):
        return float(spacing.pt)
    return float(spacing)


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    style_name = getattr(style, "name", "") or ""
    match = re.match(r"Heading\s+(\d+)$", style_name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    if style_name.lower() == "title":
        return 0
    paragraph_properties = paragraph._p.pPr
    outline_level = (
        paragraph_properties.outlineLvl if paragraph_properties is not None else None
    )
    if outline_level is not None:
        value = outline_level.get(qn("w:val"))
        if value is not None and value.isdigit():
            return int(value) + 1
    return None


def _margin_points(value: object | None) -> float | None:
    return float(value.pt) if value is not None and hasattr(value, "pt") else None


def _inline_shape_relationship_id(shape: object) -> str | None:
    inline = getattr(shape, "_inline", None)
    if inline is None:
        return None
    relationship_ids = inline.xpath(".//a:blip/@r:embed")
    return str(relationship_ids[0]) if relationship_ids else None


def parse_docx(path: Path) -> DocumentModel:
    try:
        document = load_docx(str(path))
    except (BadZipFile, PackageNotFoundError, ValueError, KeyError, OSError) as exc:
        raise _invalid_document(exc) from exc

    paragraphs: list[ParagraphModel] = []
    headings: list[HeadingModel] = []
    references: list[ReferenceModel] = []
    captions: dict[str, list[str]] = {"图": [], "表": []}
    in_references = False

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else None
        alignment = paragraph.alignment
        paragraph_model = ParagraphModel(
            index=index,
            text=text,
            style_name=style_name,
            alignment=alignment.name.lower() if alignment is not None else None,
            line_spacing=_line_spacing(paragraph),
            runs=[
                RunModel(
                    text=run.text,
                    font_name=_run_font_name(run),
                    font_size=float(run.font.size.pt)
                    if run.font.size is not None
                    else None,
                    bold=run.bold,
                    italic=run.italic,
                    underline=bool(run.underline)
                    if run.underline is not None
                    else None,
                )
                for run in paragraph.runs
            ],
        )
        paragraphs.append(paragraph_model)

        heading_level = _heading_level(paragraph)
        if heading_level is not None and text:
            headings.append(
                HeadingModel(paragraph_index=index, text=text, level=heading_level)
            )
        if REFERENCE_HEADING.match(text):
            in_references = True
            continue
        if in_references and text:
            match = REFERENCE_ENTRY.match(text)
            number = int(match.group(1) or match.group(2)) if match else None
            references.append(
                ReferenceModel(index=len(references), text=text, number=number)
            )
        caption = CAPTION_PATTERN.match(text)
        if caption:
            captions[caption.group(1)].append(text)

    tables = [
        TableModel(
            index=index,
            rows=[[cell.text.strip() for cell in row.cells] for row in table.rows],
            caption=captions["表"][index] if index < len(captions["表"]) else None,
        )
        for index, table in enumerate(document.tables)
    ]
    figures = [
        FigureModel(
            index=index,
            relationship_id=_inline_shape_relationship_id(shape),
            caption=captions["图"][index] if index < len(captions["图"]) else None,
        )
        for index, shape in enumerate(document.inline_shapes)
    ]
    sections = [
        SectionModel(
            index=index,
            top_margin=_margin_points(section.top_margin),
            right_margin=_margin_points(section.right_margin),
            bottom_margin=_margin_points(section.bottom_margin),
            left_margin=_margin_points(section.left_margin),
            header_text="\n".join(
                item.text for item in section.header.paragraphs if item.text.strip()
            ),
            footer_text="\n".join(
                item.text for item in section.footer.paragraphs if item.text.strip()
            ),
        )
        for index, section in enumerate(document.sections)
    ]
    raw_text = "\n".join(item.text for item in paragraphs if item.text)
    if not raw_text.strip():
        raise _invalid_document()
    return DocumentModel(
        file_type=".docx",
        metadata={"source": path.name},
        sections=sections,
        paragraphs=paragraphs,
        headings=headings,
        tables=tables,
        figures=figures,
        references=references,
        raw_text=raw_text,
    )
