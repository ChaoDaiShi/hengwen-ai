import base64
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def build_docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for text in paragraphs or ("衡文测试文档",):
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_structured_docx(path: Path) -> Path:
    document = DocxDocument()
    document.add_heading("衡文测试论文", level=0)
    document.add_heading("1 绪论", level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("这是使用宋体小四字号的正文。")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "规范"
    table.cell(1, 1).text = "通过"
    document.add_heading("参考文献", level=1)
    document.add_paragraph("[1] 张三. 文档审查研究[J]. 测试学报, 2025, 1(1): 1-8.")
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.header.paragraphs[0].text = "衡文测试"
    section.footer.paragraphs[0].text = "第 1 页"
    document.save(path)
    return path


def build_docx_with_inline_image(path: Path) -> Path:
    image_path = path.with_suffix(".png")
    image_path.write_bytes(ONE_PIXEL_PNG)
    document = DocxDocument()
    document.add_heading("含图片的测试论文", level=0)
    document.add_paragraph("正文内容")
    document.add_picture(str(image_path), width=Inches(1))
    document.save(path)
    return path


def build_pdf_bytes(text: str = "HengWen PDF") -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, payload in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(payload)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)
