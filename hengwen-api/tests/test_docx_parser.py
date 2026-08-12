from pathlib import Path

import pytest
from docx import Document as DocxDocument

from hengwen_api.core.exceptions import AppError, ErrorCode
from hengwen_api.document_engine.parser import parse_document
from tests.factories import build_structured_docx


def test_docx_parser_extracts_runs_headings_tables_and_sections(
    tmp_path: Path,
) -> None:
    path = build_structured_docx(tmp_path / "structured.docx")

    model = parse_document(path, ".docx")

    assert model.file_type == ".docx"
    assert model.raw_text.startswith("衡文测试论文")
    assert model.headings[0].text == "衡文测试论文"
    assert model.headings[0].level == 0
    assert model.headings[1].text == "1 绪论"
    body = next(item for item in model.paragraphs if "宋体" in item.text)
    assert body.runs[0].font_name == "宋体"
    assert body.runs[0].font_size == 12.0
    assert body.alignment == "justify"
    assert body.line_spacing == 1.5
    assert model.tables[0].rows[1] == ["规范", "通过"]
    assert model.sections[0].top_margin == pytest.approx(72.0)
    assert model.sections[0].header_text == "衡文测试"
    assert model.references[0].number == 1


def test_docx_parser_rejects_empty_document(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    DocxDocument().save(path)

    with pytest.raises(AppError) as captured:
        parse_document(path, ".docx")

    assert captured.value.code == ErrorCode.INVALID_DOCUMENT


def test_docx_parser_rejects_corrupt_package(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not-a-docx")

    with pytest.raises(AppError) as captured:
        parse_document(path, ".docx")

    assert captured.value.code == ErrorCode.INVALID_DOCUMENT
